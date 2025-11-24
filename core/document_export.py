# core/document_export.py
import pandas as pd
from docx import Document

def export_to_excel(answer_text: str, output_path: str):
    df = pd.DataFrame({"AI Response": [answer_text]})
    df.to_excel(output_path, index=False)
    return output_path


def export_to_word(answer_text: str, output_path: str):
    doc = Document()
    doc.add_heading("AI Generated Response", level=1)
    for line in answer_text.split("\n"):
        doc.add_paragraph(line)
    doc.save(output_path)
    return output_path


def export_tabulated_report(answer_text: str, output_path: str):
    doc = Document()
    doc.add_heading("AI Report (Clean Format)", level=1)

    lines = answer_text.split("\n")
    table = doc.add_table(rows=1, cols=1)
    header_cells = table.rows[0].cells
    header_cells[0].text = "AI Response"

    for line in lines:
        row_cells = table.add_row().cells
        row_cells[0].text = line.strip()

    doc.save(output_path)
    return output_path
